# -*- coding: utf-8 -*-
"""Generate a Word doc combining promo copy + screenshots for Survival Arena.
Reads promo-copy.md (plain text sections) and inserts screenshots at the right steps.# Output: 游戏推广文案-生存竞技场.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(BASE, 'screenshots')

RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x16, 0xC7, 0x9A)
GRAY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# Page setup: A4, reasonable margins
sec = doc.sections[0]
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)

# Pin document metadata to fixed timestamps so re-generation is byte-stable
# (prevents git churn from python-docx writing a fresh 'now' on every run).
import datetime as _dt
_fixed = _dt.datetime(2026, 8, 12, 0, 0, 0)
doc.core_properties.created = _fixed
doc.core_properties.modified = _fixed

def h1(text, color=RED):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = color
    r.font.name = 'Microsoft YaHei'
    return p

def h2(text, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    if color:
        r.font.color.rgb = color
    r.font.name = 'Microsoft YaHei'
    return p

def body(text, size=11, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    r.font.name = 'Microsoft YaHei'
    return p

# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = t.add_run('🎮 生存竞技场 Survival Arena')
tr.bold = True; tr.font.size = Pt(30); tr.font.color.rgb = RED; tr.font.name = 'Microsoft YaHei'
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = st.add_run('游戏推广文案 · 附配图演示')
sr.font.size = Pt(14); sr.font.color.rgb = GRAY; sr.font.name = 'Microsoft YaHei'
st.add_run('\n游戏链接：https://survival-arena.com/').bold = True

doc.add_paragraph()  # spacer

# --- 标题候选 ---
h1('📢 标题（任选其一）')
titles = [
    '1. 手机电脑都能玩的网页射击游戏，不用下载点开就玩！🔥',
    '2. 我在网页里玩到的爽快射击游戏，被好友无限救活😭',
    '3. 零安装网页射击游戏！上班摸鱼偷偷玩系列🎮',
    '4. 开局一把枪X波怪，好友一键把我捞起来继续打！🪖',
]
for x in titles:
    body(x, size=12, color=RED, bold=True)

doc.add_paragraph()

# --- 正文 ---
h1('📝 正文（含配图）')
body('一个不用下载、浏览器点开就玩的 2D 俯视角爽快射击游戏！手机⚡电脑💻都能玩，随时开一局。👇下面每一步都真实截图给你看👇')

step_sections = [
    ('#01 进入游戏', '🔑 打开就是登录界面，支持手机号/邮箱，首次登录自动注册，秒进。', '01-登录界面.png'),
    ('#02 主菜单', '🏠 进游戏先看到功能面板——开始战斗、装备升级、每日奖励、排行榜、皮肤，全在一个界面，一目了然。', '02-主菜单.png'),
    ('#03 装备与升级', '🛒 不肝也能变强！金币升属性、钻石换高级货币、星星解锁武器。4 种武器：手枪/霰弹枪/冲锋枪/狙击枪，各有独立升级树。', '03-装备与升级商店.png'),
    ('#04 每日奖励', '🎁 每天登录领好礼，连续签到 7 天送大礼包，含复活令牌！', '04-每日奖励.png'),
    ('#05 皮肤系统', '🎨 6 款皮肤随便换：默认战士/烈焰使者/暗影刺客/黄金骑士/霓虹战士/虚空领主，超帅。', '05-皮肤系统.png'),
    ('#06 排行榜', '🏆 5 大段位：青铜→白银→黄金→钻石→传说，和朋友比谁是真枪神。', '06-排行榜.png'),
    ('#07 开始战斗', '⚔️ WASD 移动 + 鼠标瞄准，自动射击，空格冲刺闪避，第 5/10/15 波打 BOSS！', '07-战斗画面WASD.png'),
    ('#08 暂停装备', '⏸️ 打一半能随时暂停换武器、买装备、升属性，战斗节奏自己掌控。', '08-暂停菜单装备.png'),
    ('#09 升级加点', '🧬 局内吃到星星就能点技能树——攻击/攻速/移速/护盾，build 随你搭。', '09-暂停菜单升级.png'),
    ('#10 战斗实况', '💥 屏幕震动 + 粒子特效 + 弹幕 BOSS，5 种敌人轮番来袭，爽到停不下来。', '10-战斗进行中.png'),
    ('#11 好友复活（压轴卖点）', '🤝 阵亡不要紧！生成专属分享链接发给好友，好友一点开，我满血复活 + 3 秒无敌护盾！关键：电脑和手机都能互救——好友在电脑、我在手机，照样救得回来！复活双方都拿 +50 金币 +5 钻石奖励，越玩越有感情~（还有看广告、用令牌、花钻石三种复活方式可选）', '11-好友复活分享.png'),
]

for caption, desc, imgfile in step_sections:
    h2(caption, GREEN)
    body(desc, size=12)
    img_path = os.path.join(IMG, imgfile)
    if os.path.exists(img_path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img_path, width=Inches(6.0))
    body('', size=4)

doc.add_paragraph()

# --- 经济 & 功能亮点 ---
h1('💰 亮点速览')
body('• 双货币经济：星星升武器 🪙 金币升属性 💎 钻石抽皮肤', size=11)
body('• 12 项成就 + 每日挑战，越玩越上瘾', size=11)
body('• 4 种武器、5 种敌人、6 款皮肤、5 大段位', size=11)
body('• 社交裂变：跨设备好友复活，电脑↔手机互救', size=11, color=RED, bold=True)
body('👉 想试试的直接点链接，不用下载：https://survival-arena.com/', size=12, bold=True)

doc.add_paragraph()

# --- 发布小贴士 ---
h1('💡 发布小贴士')
tips = [
    '小红书：正文 + 11 张图；标题用第 2 条（"被好友无限救活"最有互动点）；结尾加话题 #网页游戏 #射击游戏 #手机游戏 #不吃配置的游戏',
    '抖音/视频号：用#10 战斗实况截图当封面，短视频可循环打 BOSS 画面',
    '微博：精简成 3 张（#02 主菜单、#07 战斗、#11 复活）+ 链接',
]
for i, x in enumerate(tips, 1):
    body(f'{i}. {x}', size=11, color=GRAY)

out = os.path.join(BASE, '游戏推广文案-生存竞技场.docx')
doc.save(out)

# Rewrite the zip with a FIXED timestamp on every entry so the .docx is fully
# byte-stable (python-docx otherwise stamps each zip member with the current time,
# which makes every re-generation differ in git even when the content is unchanged).
import zipfile as _zip
import io as _io
_TS = (2026, 8, 12, 21, 0, 0)  # year,month,day,hour,min,sec for every entry
_src = open(out, 'rb').read()
_buf = _io.BytesIO()
zout = _zip.ZipFile(_buf, 'w', _zip.ZIP_DEFLATED)
zin = _zip.ZipFile(_io.BytesIO(_src))
for item in zin.infolist():
    zi = _zip.ZipInfo(item.filename, date_time=_TS)
    zi.compress_type = _zip.ZIP_DEFLATED
    zi.external_attr = item.external_attr
    zout.writestr(zi, zin.read(item.filename))
zout.close()
open(out, 'wb').write(_buf.getvalue())

print('saved:', out)
